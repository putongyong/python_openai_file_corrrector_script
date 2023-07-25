import openai
import docx
import re

# Prompt the user to input their own OpenAI API key
def get_openai_api_key():
    api_key = input("Please enter your OpenAI API key: ")
    return api_key

# Define the main function
def chat_gpt(message, api_key):
    if message.lower().startswith("read file"):
        file_path = message.split(" ")[-1]
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
     
    elif message.lower().startswith("correct file"):
        file_path = message.split(" ")[-1]
        doc = docx.Document(file_path)
        
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        chunk_size = 256
        corrected_text = ""
        
        sentences = re.split(r'(?<=[.!?])\s+', text)  # Split into sentences using regex
        
        for sentence in sentences:
            if len(sentence) <= chunk_size:
                chunk = sentence
            else:
                chunks = re.split(r'(?<=[,])\s+', sentence)  # Split by comma if necessary
                if len(chunks[0]) <= chunk_size:
                    chunk = chunks[0]
                else:
                    chunk = chunks[0] + ','
                    
            openai.api_key = api_key

            # Prompt message for correction
            prompt = f"Please correct or improve the following text:\n{chunk}"

            # Call the OpenAI API for correction
            response = openai.Completion.create(
                engine='text-davinci-003',
                prompt=prompt,
                max_tokens=128,  # Adjust as needed
                temperature=0.7,
                n=1,
                stop=None,
                timeout=None
            )
            
            # Extract the corrected text from the response
            corrected_chunk = response.choices[0].text.strip()
            print(corrected_chunk)
            corrected_text += corrected_chunk
        
        return corrected_text
        
    else:
        # Handle other input messages
        # Set up OpenAI API credentials
        openai.api_key = api_key

        response = openai.Completion.create(
            engine='text-davinci-003',
            prompt=message,
            max_tokens=128,
            temperature=0.7,
            n=1,
            stop=None,
            timeout=None
        )
        reply = response.choices[0].text.strip()
        return reply
    
api_key = get_openai_api_key()
user_input = input("You: ")

while user_input.lower() != 'exit':
    response = chat_gpt(user_input, api_key)
    print("ChatGPT:", response)
    user_input = input("You: ")
